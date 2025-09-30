from datetime import datetime
from pathlib import Path
from typing import List

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.employee import Employee
from utils.utils import get_unique_file_path
import logging

logger = logging.getLogger(__name__)


class DocumentGenerator:
    """Main class for generating kwartalny security service reports."""

    def __init__(
        self,
        firearms_file_name: Path,
        entities_file_name: Path,
        supervision_file_name: Path,
        month: str | None,
        quarter: str | None,
        start_date="",
        end_date="",
        interval="",
        employees: List[Employee] = [],
    ):
        if start_date == "" or end_date == "" or interval == "" or employees == []:
            raise Exception(
                "Brakuje którejś ze ściezek do plików źródłowych lub wynikowych"
            )

        self.start_date = start_date
        self.end_date = end_date
        self.interval = interval
        self.employees = employees
        if month is not None:
            self.month = month
            self.quarter = None
        else:
            self.quarter = quarter
            self.month = None
        self.final_df = self._load_csv(entities_file_name)
        self.firearms_df = self._load_csv(firearms_file_name)
        self.supervision_data = self._load_csv(supervision_file_name)

        # Convert date columns in firearms_df
        if "Daty dotyczące przydziału broni palnej" in self.firearms_df.columns:
            self.firearms_df["przydział"] = pd.to_datetime(
                self.firearms_df["Daty dotyczące przydziału broni palnej"],
                errors="coerce",
            )

        # Convert dates in supervision_data
        self.supervision_data["rozpoczęcie"] = pd.to_datetime(
            self.supervision_data["rozpoczęcie"], errors="coerce"
        )
        self.supervision_data["zakończenie"] = pd.to_datetime(
            self.supervision_data["zakończenie"], errors="coerce"
        )

    def _load_csv(self, file_path):
        """Helper to load CSV and handle errors."""
        try:
            df = pd.read_csv(file_path, encoding="utf-8")
            logger.info(f"Successfully loaded {file_path}")
            return df
        except FileNotFoundError:
            print(f"Error: {file_path} not found. Please ensure the file exists.")
            return pd.DataFrame()  # Return empty DataFrame to avoid further errors
        except Exception as e:
            print(f"Error loading {file_path}: {e}")
            return pd.DataFrame()

    def add_page_numbers(self, doc):
        """Dodaje numer strony w formacie '1 | Strona' z linią nad stopką."""
        section = doc.sections[0]
        footer = section.footer
        paragraph = (
            footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        )

        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p_pr = paragraph._element.get_or_add_pPr()
        p_borders = OxmlElement("w:pBdr")
        top = OxmlElement("w:top")
        top.set(qn("w:val"), "single")
        top.set(qn("w:sz"), "6")  # grubość linii
        top.set(qn("w:space"), "1")  # odstęp
        top.set(qn("w:color"), "auto")
        p_borders.append(top)
        p_pr.append(p_borders)

        run = paragraph.add_run()
        run.font.bold = True
        run.font.size = Pt(10)

        fldChar1 = OxmlElement("w:fldChar")  # begin field
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")  # polecenie
        instrText.text = "PAGE"  # type: ignore

        fldChar2 = OxmlElement("w:fldChar")  # separate field
        fldChar2.set(qn("w:fldCharType"), "separate")

        fldChar3 = OxmlElement("w:fldChar")  # end field
        fldChar3.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        separator_run = paragraph.add_run(" | ")
        separator_run.bold = True
        separator_run.font.size = Pt(10)

        text_run = paragraph.add_run("Strona")
        text_run.bold = False
        text_run.font.size = Pt(10)

    # def get_quarter_dates(self, year, quarter):
    #     """Calculate start and end dates for a given quarter."""
    #     quarters = {
    #         1: ("01-01", "03-31"),
    #         2: ("04-01", "06-30"),
    #         3: ("07-01", "09-30"),
    #         4: ("10-01", "12-31"),
    #     }
    #     start_month_day, end_month_day = quarters[quarter]

    #     quarter_start = pd.to_datetime(f"{year}-{start_month_day}")
    #     quarter_end = pd.to_datetime(f"{year}-{end_month_day}")

    #     return quarter_start, quarter_end

    def create_document_content(
        self, doc, start_date: str, end_date: str, department, current_row
    ):
        try:
            year = datetime.strptime(self.end_date, "%Y-%m-%d").year
            # quarter_start, quarter_end = self.get_quarter_dates(year, quarter)
            poz_ks = int(float(current_row["POZ KS R Umów"]))

            # Filter matching rows for Section V (Locations and Form)
            matching_rows = self.final_df[
                self.final_df["POZ KS R Umów"] == current_row["POZ KS R Umów"]
            ].copy()

            # I. Księga Realizacji
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run("I.  ")
            run.bold = True
            run.font.size = Pt(14)

            run = para.add_run(f"KSIĘGA REALIZACJI UMOWY Nr {poz_ks}")
            run.bold = True
            run.font.size = Pt(14)

            # Add contract party name
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.space_before = Pt(2)
            run = para.add_run(
                f"    Zawartej z {current_row['Oznaczenie strony lub stron umowy, z którymi przedsiębiorca zawarł umowę']}"
            )
            run.bold = True
            run.font.size = Pt(12)
            para.space_after = Pt(12)

            # II. Start date
            para = doc.add_paragraph()
            run = para.add_run(f"II. OD: {start_date}")
            run.bold = True

            # III. End date
            para = doc.add_paragraph()
            run = para.add_run(f"III. DO: {end_date}")
            run.bold = True

            # IV. Volume (Dział-CK)
            para = doc.add_paragraph()
            run = para.add_run(
                f"IV. VOL. Nr. {current_row['Dział']}-{current_row['CK']}"
            )
            run.bold = True

            # V. Location and form
            para = doc.add_paragraph()
            run = para.add_run(
                "V. Miejsce wykonywania usługi oraz forma jej wykonywania"
            )
            run.bold = True

            table = doc.add_table(rows=len(matching_rows) + 1, cols=8)
            table.style = "Table Grid"
            col_widths = [5, 8, 22, 22, 18, 9, 9, 7]

            for row_obj in table.rows:
                for idx, cell in enumerate(row_obj.cells):
                    cell.width = Inches(col_widths[idx] / 100 * 7.5)

            headers = [
                "L.p.",
                "Księga",
                "Określenie obiektu",
                "Adres Obiektu",
                "Forma Wykonywanej Usługi",
                "Data rozpoczęcia",
                "Data zakończenia",
                "Uwagi",
            ]
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True

            for idx, row_data in matching_rows.iterrows():
                row_num = idx - matching_rows.index[0] + 1
                data_row = table.rows[row_num]

                data_row.cells[0].text = f"{row_num}."
                data_row.cells[1].text = str(poz_ks)
                data_row.cells[2].text = (
                    str(row_data["Określenie obiektu"])
                    if pd.notna(row_data["Określenie obiektu"])
                    else ""
                )
                data_row.cells[3].text = (
                    str(row_data["Adres Obiektu"])
                    if pd.notna(row_data["Adres Obiektu"])
                    else ""
                )
                data_row.cells[4].text = (
                    str(row_data["Forma wykonywanej usługi"])
                    if pd.notna(row_data["Forma wykonywanej usługi"])
                    else ""
                )
                data_row.cells[5].text = (
                    pd.to_datetime(row_data["Data rozpoczęcia usługi"]).strftime(
                        "%Y-%m-%d"
                    )
                    if pd.notna(row_data["Data rozpoczęcia usługi"])
                    else ""
                )
                data_row.cells[6].text = (
                    pd.to_datetime(row_data["Data zakończenia usługi"]).strftime(
                        "%Y-%m-%d"
                    )
                    if pd.notna(row_data["Data zakończenia usługi"])
                    else ""
                )
                data_row.cells[7].text = (
                    str(row_data["Uwagi"]) if pd.notna(row_data["Uwagi"]) else ""
                )

            # VI. PRACOWNICY OCHRONY WYKONUJĄCY USŁUGĘ
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run("VI. PRACOWNICY OCHRONY WYKONUJĄCY USŁUGĘ")
            run.bold = True

            current_ck = str(current_row["CK"]).strip()

            filtered_employees = [e for e in self.employees if e.ck == current_ck]

            rows_needed = max(
                2, len(filtered_employees) + 3
            )  # Ensure at least 2 rows for header + 1 empty for appearance
            table = doc.add_table(rows=rows_needed, cols=8)
            table.style = "Table Grid"

            col_widths = [5, 15, 15, 12, 23, 12, 12, 6]
            for row_obj in table.rows:
                for idx, cell in enumerate(row_obj.cells):
                    cell.width = Inches(col_widths[idx] / 100 * 7.5)

            headers = [
                "L.p.",
                "Nazwisko",
                "Imię",
                "Numer Legitymacji",
                "Funkcja w obiekcie",
                "Data rozpoczęcia",
                "Data zakończenia",
                "Uwagi",
            ]
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True

            if len(filtered_employees) == 0:
                contract_name = current_row[
                    "Oznaczenie strony lub stron umowy, z którymi przedsiębiorca zawarł umowę"
                ]
                logger.warning(
                    f"\nWARNING: No employees found for contract: {contract_name} (CK: {current_ck}, {start_date} - {end_date})"
                )

            logger.info(
                f"\nProcessing contract: {current_row['Oznaczenie strony lub stron umowy, z którymi przedsiębiorca zawarł umowę']}"
            )
            logger.info(f"CK value: {current_ck}")
            logger.info(f"Date range: {start_date} to {end_date}")
            logger.info(f"Found {len(filtered_employees)} matching employees")
            if not len(filtered_employees) > 0:
                logger.info("First few matching employees:")
                logger.info(filtered_employees[:3])

            for idx in range(min(len(filtered_employees), rows_needed - 1)):
                data_row = table.rows[idx + 1]
                emp = filtered_employees[idx]

                data_row.cells[0].text = f"{idx + 1}."
                data_row.cells[1].text = str(emp.last_name)
                data_row.cells[2].text = str(emp.first_name)
                data_row.cells[3].text = str(emp.kod)
                data_row.cells[4].text = str(emp.position)
                data_row.cells[5].text = str(start_date)
                data_row.cells[6].text = str(emp.release_date)
                data_row.cells[7].text = ""

            for idx in range(
                len(filtered_employees), rows_needed - 1
            ):  # Fill remaining rows with empty strings
                data_row = table.rows[idx + 1]
                for cell in data_row.cells:
                    cell.text = ""

            # VII. PRACOWNICY OCHRONY SPRAWUJĄCY NADZÓR
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run(
                "VII. PRACOWNICY OCHRONY SPRAWUJĄCY NADZÓR NAD PRACOWNIKAMI OCHRONY WYKONUJĄCYMI USŁUGĘ"
            )
            run.bold = True

            dep_data = self.supervision_data[
                self.supervision_data["Dział"] == department
            ].copy()
            filtered_data = dep_data[
                (dep_data["rozpoczęcie"] <= pd.to_datetime(end_date))
                & (
                    (dep_data["zakończenie"].isna())
                    | (dep_data["zakończenie"] >= pd.to_datetime(start_date))
                )
            ]

            rows_needed_supervision = max(3, len(filtered_data) + 1)
            table = doc.add_table(rows=rows_needed_supervision, cols=8)
            table.style = "Table Grid"
            col_widths = [5, 15, 15, 12, 23, 12, 12, 6]

            for row_obj in table.rows:
                for idx, cell in enumerate(row_obj.cells):
                    cell.width = Inches(col_widths[idx] / 100 * 7.5)

            headers = [
                "L.p.",
                "Nazwisko",
                "Imię",
                "Numer Legitymacji",
                "Funkcja w obiekcie",
                "Daty rozpoczęcie",
                "Daty zakończenie",
                "Uwagi",
            ]
            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True

            for idx in range(min(len(filtered_data), rows_needed_supervision - 1)):
                data_row = table.rows[idx + 1]
                row_s = filtered_data.iloc[idx]

                start_date_s = ""
                if pd.notna(row_s["rozpoczęcie"]):
                    # Use the later of the quarter start or the actual start date
                    start_date_s = max(
                        row_s["rozpoczęcie"], pd.to_datetime(start_date)
                    ).strftime("%Y-%m-%d")

                end_date_s = ""
                if pd.notna(row_s["zakończenie"]):
                    end_date_s = row_s["zakończenie"].strftime("%Y-%m-%d")
                else:
                    # If no end date, use empty string because it didn't end
                    end_date_s = ""

                data_row.cells[0].text = f"{idx + 1}."
                data_row.cells[1].text = str(row_s["Nazwisko"])
                data_row.cells[2].text = str(row_s["Imię"])
                data_row.cells[3].text = str(row_s["Nr legitymacji"]).strip('"')
                data_row.cells[4].text = str(row_s["Funkcja w obiekcie"])
                data_row.cells[5].text = start_date_s
                data_row.cells[6].text = end_date_s
                data_row.cells[7].text = (
                    str(row_s["Uwagi"]) if pd.notna(row_s["Uwagi"]) else ""
                )

            for idx in range(
                len(filtered_data), rows_needed_supervision - 1
            ):  # Fill remaining with empty strings
                data_row = table.rows[idx + 1]
                for cell in data_row.cells:
                    cell.text = ""

            # VIII. ILOŚĆ I RODZAJ BRONI PALNEJ
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run(
                "VIII. ILOŚĆ I RODZAJ BRONI PALNEJ PRZYDZIELONEJ PRACOWNIKOM OCHRONY DO WYKONANIA USŁUGI"
            )
            run.bold = True

            table = doc.add_table(rows=2, cols=9)
            table.style = "Table Grid"
            col_widths = [5, 12, 10, 8, 8, 27, 12, 12, 6]

            for row_obj in table.rows:
                for idx, cell in enumerate(row_obj.cells):
                    cell.width = Inches(col_widths[idx] / 100 * 7.5)

            headers = [
                "L.p.",
                "Rodzaj broni palnej",
                "Marka broni",
                "Kaliber",
                "Ilość",
                "Obiekt, do którego przydzielono pracownikom broń palną",
                "Przydział broni palnej",
                "Cofnięcie przydziału broni palnej",
                "Uwagi",
            ]

            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True

            data_row_firearms = table.rows[1]
            data_cells_firearms = ["1", "Nie przyznano", "", "", "", "", "", "", ""]

            if department == "MON":
                mon_firearms = self.firearms_df[
                    (
                        self.firearms_df[
                            "Obiekt, do którego przydzielono pracownikom broń palną"
                        ]
                        == "MON"
                    )
                    & (self.firearms_df["przydział"].notna())
                ].copy()

                if not mon_firearms.empty:
                    row_f = mon_firearms.iloc[0]

                    assignment_date_f = ""
                    if pd.notna(row_f["przydział"]):
                        # If assignment date is before quarter start, use quarter start
                        assignment_date_f = max(
                            row_f["przydział"], pd.to_datetime(start_date)
                        ).strftime("%Y-%m-%d")

                    assignment_end_date_f = ""
                    # fix csv file
                    # if pd.notna(row_f["cofnięcie przydziału"]):
                    #     assignment_date_f = max(
                    #         row_f["cofnięcie przydziału"], pd.to_datetime(end_date)
                    #     ).strftime("%Y-%m-%d")

                    data_cells_firearms = [
                        "1",
                        (
                            str(row_f["Rodzaj broni palnej"])
                            if pd.notna(row_f["Rodzaj broni palnej"])
                            else ""
                        ),
                        (
                            str(row_f["Marka broni"])
                            if pd.notna(row_f["Marka broni"])
                            else ""
                        ),
                        str(row_f["Kaliber"]) if pd.notna(row_f["Kaliber"]) else "",
                        str(row_f["Ilość"]) if pd.notna(row_f["Ilość"]) else "",
                        (
                            str(
                                row_f[
                                    "Obiekt, do którego przydzielono pracownikom broń palną"
                                ]
                            )
                            if pd.notna(
                                row_f[
                                    "Obiekt, do którego przydzielono pracownikom broń palną"
                                ]
                            )
                            else ""
                        ),
                        assignment_date_f,
                        assignment_end_date_f,
                        str(row_f["Uwagi"]) if pd.notna(row_f["Uwagi"]) else "",
                    ]

            for col, value in enumerate(data_cells_firearms):
                data_row_firearms.cells[col].text = value

            # IX. ILOŚĆ I RODZAJ ŚRODKÓW PRZYMUSU BEZPOŚREDNIEGO
            doc.add_paragraph()
            para = doc.add_paragraph()
            run = para.add_run(
                "IX. ILOŚĆ I RODZAJ ŚRODKÓW PRZYMUSU BEZPOŚREDNIEGO PRZYDZIELONYCH PRACOWNIKOM OCHRONY DO WYKONANIA USŁUGI"
            )
            run.bold = True

            table = doc.add_table(rows=2, cols=6)
            table.style = "Table Grid"
            col_widths = [5, 35, 10, 25, 15, 10]

            for row_obj in table.rows:
                for idx, cell in enumerate(row_obj.cells):
                    cell.width = Inches(col_widths[idx] / 100 * 7.5)

            headers = [
                "L.p.",
                "Rodzaj środka przymusu bezpośredniego",
                "Ilość",
                "Obiekt do którego przydzielono pracownikom ś.p.b.",
                "Daty przydziału",
                "Uwagi",
            ]

            for i, header in enumerate(headers):
                cell = table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True

            data_row_coercion = table.rows[1]
            data_cells_coercion = ["1", "Nie przyznano", "", "", "", ""]

            for col, value in enumerate(data_cells_coercion):
                data_row_coercion.cells[col].text = value

        except Exception as e:
            print(
                f"Error creating document content for POZ KS R Umów {current_row['POZ KS R Umów']}: {e}"
            )
            raise

    def generate_quarterly_reports(
        self, output_path: Path, start_date: str, end_date: str
    ):
        """Generates kwartalny documents with all records."""
        try:

            doc = Document()

            # Set document margins
            for section in doc.sections:
                section.left_margin = Inches(0.5)
                section.right_margin = Inches(0.5)
                section.top_margin = Inches(0.5)
                section.bottom_margin = Inches(1)

            # Add page numbers
            self.add_page_numbers(doc)  # Correctly call as a method

            # Ensure 'POZ KS R Umów' and 'Dział' are treated consistently
            filtered_contracts = self.final_df[
                self.final_df["POZ KS R Umów"].notna()
                & (self.final_df["POZ KS R Umów"].astype(str) != "")
            ].copy()

            for dept in ["MON", "OFS"]:
                dept_rows = filtered_contracts[
                    filtered_contracts["Dział"] == dept
                ].copy()

                for idx, row in dept_rows.iterrows():
                    contract_start_date = pd.to_datetime(
                        row["Data rozpoczęcia usługi"], errors="coerce"
                    )
                    contract_end_date = pd.to_datetime(
                        row["Data zakończenia usługi"], errors="coerce"
                    )

                    # Filter contracts based on their start and end dates relative to the quarter
                    # A contract is relevant if it started by the end of the quarter AND (it has no end date OR it ends after the quarter started)
                    is_active_in_quarter = (
                        pd.notna(contract_start_date)
                        and contract_start_date <= pd.to_datetime(end_date)
                    ) and (
                        pd.isna(contract_end_date)
                        or contract_end_date >= pd.to_datetime(start_date)
                    )

                    if is_active_in_quarter:
                        try:
                            if (
                                doc.paragraphs and len(doc.paragraphs) > 1
                            ):  # Add page break only if there's existing content
                                doc.add_page_break()

                            self.create_document_content(
                                doc, start_date, end_date, dept, row
                            )

                        except Exception as e:
                            print(
                                f"Error processing document for {dept} - POZ KS R Umów {row['POZ KS R Umów']}: {e}"
                            )

            file_name = f"Raport_{self.start_date}_{self.end_date}.docx"
            file_path = output_path / file_name
            file_path = get_unique_file_path(file_path)
            doc.save(str(file_path))
            doc.save(str(file_path))
            print(f"Created document: {file_path}")

        except Exception as e:
            print(f"Error in document generation process: {e}")

    def create_folder_structure(self):
        output_path = Path(__file__) / "results"

        if output_path.exists() == False:
            output_path.mkdir(parents=True, exist_ok=True)
            return self.create_folder_structure()

        start_dt = datetime.strptime(self.start_date, "%Y-%m-%d")
        end_dt = datetime.strptime(self.end_date, "%Y-%m-%d")
        if start_dt.year != end_dt.year:
            raise Exception("Year of start date and end date does not match")

        year_path = output_path / str(start_dt.year)

        if not year_path.exists():
            year_path.mkdir(parents=True, exist_ok=True)
            logger.info(f"Folder '{year_path}' created.")
        else:
            logger.info(f"Folder '{year_path}' already exists.")

        if self.interval == "miesieczny" and self.month != None:
            month_folder = year_path / f"{start_dt.month}"
            if not month_folder.exists():
                month_folder.mkdir()
                logger.info(f"Folder '{month_folder}' created.")
            else:
                logger.info(f"Folder '{month_folder}' already exists.")

            return month_folder

        elif self.interval == "kwartalny" and self.quarter != None:
            quarter_folder = year_path / f"{self.quarter}"
            if not quarter_folder.exists():
                quarter_folder.mkdir()
                logger.info(f"Folder '{quarter_folder}' created.")
            else:
                logger.info(f"Folder '{quarter_folder}' already exists.")

            return quarter_folder

        else:
            logger.error(
                "No valid interval (miesieczny/kwartalny) or missing variable."
            )
            raise Exception(
                "No valid interval (miesieczny/kwartalny) or missing variable."
            )
